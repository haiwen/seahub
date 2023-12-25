# Copyright (c) 2012-2016 Seafile Ltd.
import os
import io
import re
import jwt
import docx
import time
import logging
import requests
import datetime
import urllib.request
import urllib.parse
import urllib.error

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_COLOR_INDEX

from django.urls import reverse
from urllib.parse import urljoin
from rest_framework import status

from seaserv import ccnet_api, seafile_api

from seahub.api2.utils import api_error
from seahub.tags.models import FileUUIDMap
from seahub.base.templatetags.seahub_tags import email2nickname, email2contact_email
from seahub.utils import get_log_events_by_time, is_pro_version, is_org_context, \
        gen_inner_file_get_url, get_service_url

from seahub.settings import SEADOC_PRIVATE_KEY, FILE_CONVERTER_SERVER_URL

try:
    from seahub.settings import MULTI_TENANCY
except ImportError:
    MULTI_TENANCY = False

logger = logging.getLogger(__name__)


def api_check_group(func):
    """
    Decorator for check if group valid
    """
    def _decorated(view, request, group_id, *args, **kwargs):
        group_id = int(group_id)  # Checked by URL Conf
        try:
            group = ccnet_api.get_group(int(group_id))
        except Exception as e:
            logger.error(e)
            error_msg = 'Internal Server Error'
            return api_error(status.HTTP_500_INTERNAL_SERVER_ERROR, error_msg)

        if not group:
            error_msg = 'Group %d not found.' % group_id
            return api_error(status.HTTP_404_NOT_FOUND, error_msg)

        return func(view, request, group_id, *args, **kwargs)

    return _decorated


def add_org_context(func):
    def _decorated(view, request, *args, **kwargs):
        if is_org_context(request):
            org_id = request.user.org.org_id
        else:
            org_id = None
        return func(view, request, org_id=org_id, *args, **kwargs)

    return _decorated


def is_org_user(username, org_id=None):
    """ Check if an user is an org user.

    Keyword arguments:
    org_id -- An integer greater than zero. If provided,
    check if the user is a member of the specific org.
    """

    if not is_pro_version() or not MULTI_TENANCY:
        return False

    try:
        if org_id:
            # Return non-zero if True, otherwise 0.
            return ccnet_api.org_user_exists(org_id, username) != 0
        else:
            orgs = ccnet_api.get_orgs_by_user(username)
            return len(orgs) > 0
    except Exception as e:
        logger.error(e)
        return False


def get_user_contact_email_dict(email_list):
    email_list = set(email_list)
    user_contact_email_dict = {}
    for email in email_list:
        if email not in user_contact_email_dict:
            user_contact_email_dict[email] = email2contact_email(email)

    return user_contact_email_dict


def get_user_name_dict(email_list):
    email_list = set(email_list)
    user_name_dict = {}
    for email in email_list:
        if email not in user_name_dict:
            user_name_dict[email] = email2nickname(email)

    return user_name_dict


def get_repo_dict(repo_id_list):
    repo_id_list = set(repo_id_list)
    repo_dict = {}
    for repo_id in repo_id_list:
        if repo_id not in repo_dict:
            repo_dict[repo_id] = ''
            repo = seafile_api.get_repo(repo_id)
            if repo:
                repo_dict[repo_id] = repo

    return repo_dict


def get_group_dict(group_id_list):
    group_id_list = set(group_id_list)
    group_dict = {}
    for group_id in group_id_list:
        if group_id not in group_dict:
            group_dict[group_id] = ''
            group = ccnet_api.get_group(int(group_id))
            if group:
                group_dict[group_id] = group

    return group_dict


def check_time_period_valid(start, end):
    if not start or not end:
        return False

    # check the date format, should be like '2015-10-10'
    date_re = re.compile(r'^(\d{4})\-([1-9]|0[1-9]|1[012])\-([1-9]|0[1-9]|[12]\d|3[01])$')
    if not date_re.match(start) or not date_re.match(end):
        return False

    return True


def get_log_events_by_type_and_time(log_type, start, end):
    start_struct_time = datetime.datetime.strptime(start, "%Y-%m-%d")
    start_timestamp = time.mktime(start_struct_time.timetuple())

    end_struct_time = datetime.datetime.strptime(end, "%Y-%m-%d")
    end_timestamp = time.mktime(end_struct_time.timetuple())
    end_timestamp += 24 * 60 * 60

    events = get_log_events_by_time(log_type, start_timestamp, end_timestamp)
    events = events if events else []
    return events


def generate_links_header_for_paginator(base_url, page, per_page, total_count, option_dict={}):

    def is_first_page(page):
        return True if page == 1 else False

    def is_last_page(page, per_page, total_count):
        if page * per_page >= total_count:
            return True
        else:
            return False

    if not isinstance(option_dict, dict):
        return ''

    query_dict = {'page': 1, 'per_page': per_page}
    query_dict.update(option_dict)

    # generate first page url
    first_page_url = base_url + '?' + urllib.parse.urlencode(query_dict)

    # generate last page url
    last_page_query_dict = {'page': (total_count / per_page) + 1}
    query_dict.update(last_page_query_dict)
    last_page_url = base_url + '?' + urllib.parse.urlencode(query_dict)

    # generate next page url
    next_page_query_dict = {'page': page + 1}
    query_dict.update(next_page_query_dict)
    next_page_url = base_url + '?' + urllib.parse.urlencode(query_dict)

    # generate prev page url
    prev_page_query_dict = {'page': page - 1}
    query_dict.update(prev_page_query_dict)
    prev_page_url = base_url + '?' + urllib.parse.urlencode(query_dict)

    # generate `Links` header
    links_header = ''
    if is_first_page(page):
        links_header = '<%s>; rel="next", <%s>; rel="last"' % (next_page_url, last_page_url)
    elif is_last_page(page, per_page, total_count):
        links_header = '<%s>; rel="first", <%s>; rel="prev"' % (first_page_url, prev_page_url)
    else:
        links_header = '<%s>; rel="next", <%s>; rel="last", <%s>; rel="first", <%s>; rel="prev"' % \
                (next_page_url, last_page_url, first_page_url, prev_page_url)

    return links_header


def get_user_quota_usage_and_total(email, org_id=''):
    try:
        if org_id:
            quota_usage = seafile_api.get_org_user_quota_usage(org_id, email)
            quota_total = seafile_api.get_org_user_quota(org_id, email)
        else:
            orgs = ccnet_api.get_orgs_by_user(email)
            if orgs:
                org_id = orgs[0].org_id
                quota_usage = seafile_api.get_org_user_quota_usage(org_id, email)
                quota_total = seafile_api.get_org_user_quota(org_id, email)
            else:
                quota_usage = seafile_api.get_user_self_usage(email)
                quota_total = seafile_api.get_user_quota(email)
    except Exception as e:
        logger.error(e)
        quota_usage = -1
        quota_total = -1
    return quota_usage, quota_total


def convert_file_gen_headers():
    payload = {'exp': int(time.time()) + 300, }
    token = jwt.encode(payload, SEADOC_PRIVATE_KEY, algorithm='HS256')
    return {"Authorization": "Token %s" % token}


def convert_file(path, username, doc_uuid, download_token, upload_token, src_type, dst_type):
    headers = convert_file_gen_headers()
    params = {
        'path': path,
        'username': username,
        'doc_uuid': doc_uuid,
        'download_token': download_token,
        'upload_token': upload_token,
        'src_type': src_type,
        'dst_type': dst_type,
    }
    url = urljoin(FILE_CONVERTER_SERVER_URL, '/api/v1/file-convert/')
    resp = requests.post(url, json=params, headers=headers, timeout=30)

    return resp


def sdoc_convert_to_docx(path, username, doc_uuid, download_token,
                         upload_token, src_type, dst_type):

    headers = convert_file_gen_headers()
    params = {
        'path': path,
        'username': username,
        'doc_uuid': doc_uuid,
        'download_token': download_token,
        'upload_token': upload_token,
        'src_type': src_type,
        'dst_type': dst_type,
    }
    url = urljoin(FILE_CONVERTER_SERVER_URL, '/api/v1/sdoc-convert-to-docx/')
    resp = requests.post(url, json=params, headers=headers, timeout=30)

    return resp


def sdoc_export_to_docx(path, username, doc_uuid, download_token,
                        src_type, dst_type):

    headers = convert_file_gen_headers()
    params = {
        'path': path,
        'username': username,
        'doc_uuid': doc_uuid,
        'download_token': download_token,
        'src_type': src_type,
        'dst_type': dst_type,
    }
    url = urljoin(FILE_CONVERTER_SERVER_URL, '/api/v1/sdoc-export-to-docx/')
    resp = requests.post(url, json=params, headers=headers, timeout=30)

    return resp


def convert_sdoc_to_docx(file_content_json, file_uuid, username):

    def add_hyperlink(paragraph, url, text, color):
        """
        A function that places a hyperlink within a paragraph object.

        :param paragraph: The paragraph we are adding the hyperlink to.
        :param url: A string containing the required url
        :param text: The text displayed for the url
        :return: The hyperlink object
        """

        # This gets access to the document.xml.rels file and gets a new relation id value
        part = paragraph.part
        r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

        # Create the w:hyperlink tag and add needed values
        hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
        hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

        # Create a w:r element
        new_run = docx.oxml.shared.OxmlElement('w:r')

        # Create a new w:rPr element
        rPr = docx.oxml.shared.OxmlElement('w:rPr')

        # Add color if it is given
        if color:
            c = docx.oxml.shared.OxmlElement('w:color')
            c.set(docx.oxml.shared.qn('w:val'), color)
            rPr.append(c)

        # Join all the xml elements together add add the required text to the w:r element
        new_run.append(rPr)
        new_run.text = text
        hyperlink.append(new_run)

        paragraph._p.append(hyperlink)

        return hyperlink

    def extract_text_in_table_recursively(data):

        text_list = []
        if isinstance(data, list):
            for item in data:
                text_list.extend(extract_text_in_table_recursively(item))
        elif isinstance(data, dict):
            if 'text' in data:
                text_list.append(data['text'])
            for key, value in data.items():
                text_list.extend(extract_text_in_table_recursively(value))

        return text_list

    def search_sdoc_node_recursively(children_list, type_sq=[], top_type=''):

        if 'text' in children_list[0]:
            if top_type == "ordered_list" and type_sq.count('ordered_list') == 1:
                type_content_list.append(['ordered_list_2', children_list])
            elif top_type == "ordered_list" and type_sq.count('ordered_list') >= 2:
                type_content_list.append(['ordered_list_3', children_list])
            elif top_type == "unordered_list" and type_sq.count('unordered_list') == 1:
                type_content_list.append(['ordered_list_2', children_list])
            elif top_type == "unordered_list" and type_sq.count('unordered_list') >= 2:
                type_content_list.append(['ordered_list_3', children_list])
            else:
                type_content_list.append([top_type, children_list])
        else:
            if top_type == 'table':
                table_text_list = extract_text_in_table_recursively(children_list)
                sub_length = len(children_list[0]['children'])
                new_table_text_list = []
                for i in range(0, len(table_text_list), sub_length):
                    new_table_text_list.append(table_text_list[i:i + sub_length])

                type_content_list.append([top_type, new_table_text_list])
            else:
                for children in children_list:
                    current_type = children.get('type', 'no type')
                    sub_children_list = children.get('children', [])
                    search_sdoc_node_recursively(sub_children_list,
                                                 type_sq + [current_type],
                                                 top_type=top_type)

    sdoc_node_list = file_content_json.get('children', [])
    type_content_list = []
    for sdoc_node in sdoc_node_list:
        top_sdoc_type = sdoc_node.get('type', '')
        children_list = sdoc_node.get('children', '')
        search_sdoc_node_recursively(children_list, top_type=top_sdoc_type)

    document = Document()

    for type_content in type_content_list:

        sdoc_type = type_content[0]
        content = type_content[1]

        if sdoc_type == 'title':
            docx_paragraph = document.add_heading(level=0)
        if sdoc_type == 'subtitle':
            docx_paragraph = document.add_paragraph(style="Subtitle")
        if sdoc_type == 'header1':
            docx_paragraph = document.add_heading(level=1)
        if sdoc_type == 'header2':
            docx_paragraph = document.add_heading(level=2)
        if sdoc_type == 'header3':
            docx_paragraph = document.add_heading(level=3)
        if sdoc_type == 'header4':
            docx_paragraph = document.add_heading(level=4)
        if sdoc_type == 'header5':
            docx_paragraph = document.add_heading(level=5)
        if sdoc_type == 'header6':
            docx_paragraph = document.add_heading(level=6)
        if sdoc_type == 'paragraph':
            docx_paragraph = document.add_paragraph()
        if sdoc_type == 'blockquote':
            docx_paragraph = document.add_paragraph(style="Intense Quote")
        if sdoc_type == 'ordered_list':
            docx_paragraph = document.add_paragraph(style="List Number")
        if sdoc_type == 'ordered_list_2':
            docx_paragraph = document.add_paragraph(style="List Number 2")
        if sdoc_type == 'ordered_list_3':
            docx_paragraph = document.add_paragraph(style="List Number 3")
        if sdoc_type in ('unordered_list', 'check_list_item'):
            docx_paragraph = document.add_paragraph(style="List Bullet")
        if sdoc_type == 'unordered_list_2':
            docx_paragraph = document.add_paragraph(style="List Bullet 2")
        if sdoc_type == 'unordered_list_3':
            docx_paragraph = document.add_paragraph(style="List Bullet 3")

        if sdoc_type == 'code_block':

            docx_paragraph = document.add_paragraph(style="No Spacing")
            docx_paragraph.paragraph_format.left_indent = Inches(0.2)

            for text_dict in content:
                text = text_dict.get('text', '')
                run = docx_paragraph.add_run(text)
                run.font.size = Pt(10)
                run.font.name = 'Courier New'

        elif sdoc_type == 'paragraph' and \
                any(item.get('type') == 'link' for item in content):

            # add hyperlink to docx

            # ['paragraph',
            #   [{'id': 'TQdHtyxhQfm8ipm76cVKKg', 'text': ''},
            #    {'children': [{'id': 'VFGENWpbTNeMRb-16QgdNA',
            #                   'text': '127.0.0.1 link title'}],
            #     'href': 'http://127.0.0.1/link-address/',
            #     'id': 'Co9L-c-SQmWk4yxHSXu5tg',
            #     'title': '127.0.0.1 link title',
            #     'type': 'link'},
            #    {'id': 'Pwqf3nbSTWmIFbwrFo1Eow', 'text': ''}]],

            link_href = ''
            link_title = ''
            for item in content:
                if 'href' in item:
                    link_href = item['href']
                if 'title' in item:
                    link_title = item['title']

            docx_paragraph = document.add_paragraph()
            add_hyperlink(docx_paragraph, link_href, link_title, "0000FF")

        elif sdoc_type == 'paragraph' and \
                any(item.get('type') in ('sdoc_link', 'file_link') for item in content):

            # add sdoc/file link to docx

            # ['paragraph',
            #  [{'id': 'D8omdcCLR4eLB3o4f0yOxw', 'text': ' '},
            #   {'children': [{'id': 'KFM5z7zvTaOcZyaT1zBhHQ', 'text': '987.sdoc'}],
            #    'display_type': 'icon_link',
            #    'doc_uuid': '45b266e4-17a5-475d-b601-10aa8001ea80',
            #    'id': 'bIwxx0mMQVKRFo3LlYwf6A',
            #    'title': '987.sdoc',
            #    'type': 'sdoc_link'},
            #   {'id': 'G5WmlQ4tSpO4IH5CDFCdUA', 'text': ' '}]],

            doc_uuid = ''
            doc_title = ''
            for item in content:
                if 'doc_uuid' in item:
                    doc_uuid = item['doc_uuid']
                if 'title' in item:
                    doc_title = item['title']
            doc_url = get_service_url() + reverse('seadoc_file_view', args=[doc_uuid])
            docx_paragraph = document.add_paragraph()
            add_hyperlink(docx_paragraph, doc_url, doc_title, "0000FF")

        elif sdoc_type in ('paragraph', 'image_block') and \
                any(item.get('type') == 'image' for item in content):

            # add image to docx

            # ['paragraph',
            #  [{'id': 'VL579VQRQdOjJCKkjRXXNA', 'text': ''},
            #   {'children': [{'id': 'dp7gIr5aSEa6GtK3-vi68g', 'text': ''}],
            #    'data': {'src': '/image-1702627227876.png'},
            #    'id': 'TEPevi-FQo-unZRBSlnd3A',
            #    'type': 'image'},
            #   {'id': 'SQjLfnvBSimn695OZtyGnw', 'text': ''}]],

            image_file_path = ''
            for item in content:
                if 'data' in item:
                    image_file_path = item['data']['src']

            uuid_map = FileUUIDMap.objects.get_fileuuidmap_by_uuid(file_uuid)
            repo_id = uuid_map.repo_id
            image_file_path = f"/images/sdoc/{file_uuid}{image_file_path}"
            image_file_id = seafile_api.get_file_id_by_path(repo_id,
                                                            image_file_path)
            download_token = seafile_api.get_fileserver_access_token(repo_id, image_file_id,
                                                                     'download', username,
                                                                     use_onetime=False)
            image_file_name = os.path.basename(image_file_path)
            image_url = gen_inner_file_get_url(download_token, image_file_name)
            resp = requests.get(image_url)
            image_content = resp.content
            document.add_picture(io.BytesIO(image_content), width=Inches(5))

        elif sdoc_type == 'table':

            # add table to docx

            # ['table', [['1', '2', '3', '4'], ['a', 'b', 'c', 'd']]]

            table = document.add_table(rows=len(content), cols=len(content[0]))

            def fulfill_table(table, content):
                for i, row in enumerate(content):
                    for j, value in enumerate(row):
                        table.cell(i, j).text = value

            fulfill_table(table, content)

        elif sdoc_type == 'callout':

            docx_paragraph = document.add_paragraph()
            for text_dict in content:
                text = text_dict.get('text', '')
                run = docx_paragraph.add_run(text)
                run.font.highlight_color = WD_COLOR_INDEX.GRAY_25

        else:

            for text_dict in content:

                text = text_dict.get('text', '') or text_dict.get('href', '')
                run = docx_paragraph.add_run(text)

                bold = text_dict.get('bold', False)
                run.bold = True if bold else False

                italic = text_dict.get('italic', False)
                run.italic = True if italic else False

    memory_stream = io.BytesIO()
    document.save(memory_stream)
    docx_content = memory_stream.getvalue()
    return docx_content
